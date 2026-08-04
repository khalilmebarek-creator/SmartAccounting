# -*- coding: utf-8 -*-
"""Generate thesis_template.docx for a Master's thesis (Algerian convention)."""
import docx
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DARK = RGBColor(0x00, 0x00, 0x00)


def set_run_font(run, size=12, bold=False, italic=False, complex_thread_font="Times New Roman"):
    """Force TNR (incl. complex scripts e.g. Arabic) + name courier excluded."""
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = DARK
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), complex_thread_font)
    rFonts.set(qn("w:cs"), complex_thread_font)


def para(doc, text="", size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_after=6, line=1.5, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = line
    if indent:
        pf.first_line_indent = Cm(indent)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def heading(doc, text, level=1, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(18 if level == 1 else 12)
    pf.space_after = Pt(10 if level == 1 else 6)
    pf.keep_with_next = True
    sizes = {1: 14, 2: 12, 3: 12}
    ital = level == 3
    r = p.add_run(text)
    set_run_font(r, size=sizes.get(level, 12), bold=True, italic=ital)
    # make it a real outline level so TOC field picks it up
    pPr = p._p.get_or_add_pPr()
    ol = OxmlElement("w:outlineLvl")
    ol.set(qn("w:val"), str(level - 1))
    pPr.append(ol)
    return p


def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_page_number_field(paragraph):
    """Insert { PAGE } field into a footer paragraph."""
    run = paragraph.add_run()
    set_run_font(run, size=10)
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)


def set_page_number_format(section, fmt, start):
    """Set w:pgNumType on the section's sectPr (e.g. lowerRoman / decimal, start value)."""
    sectPr = section._sectPr
    pg = sectPr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sectPr.append(pg)
    pg.set(qn("w:fmt"), fmt)
    pg.set(qn("w:start"), str(start))


def setup_page(doc, section):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)


def tfs(doc):
    """Times New Roman default style."""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    style.paragraph_format.line_spacing = 1.5


def set_cell(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)


def build():
    doc = Document()
    tfs(doc)

    # ---- Front matter section (roman numerals) ----
    sec = doc.sections[0]
    setup_page(doc, sec)
    footer = sec.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(fp)
    set_page_number_format(sec, "lowerRoman", 1)

    # ========== PAGE DE GARDE ==========
    para(doc, "الجمهورية الجزائرية الديمقراطية الشعبية", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.2)
    para(doc, "République Algérienne Démocratique et Populaire", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.2)
    para(doc, "Ministère de l'Enseignement Supérieur et de la Recherche Scientifique", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.2)
    para(doc, "[Nom de l'Université] — [Faculté] — [Département]", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.2)
    para(doc, "", space_after=24)
    para(doc, "Mémoire de Master", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Filière: [Filière] — Spécialité: [Spécialité]", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "", space_after=18)
    para(doc, "SMART ACCOUNTING PLATFORM", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "Design and Implementation of a Tax-Compliant Financial Analysis System for Algerian SMEs", size=13, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "", space_after=24)
    # Placeholder innovation / 1275 note
    para(doc, "Réalisé dans le cadre du mécanisme « un diplôme, une startup / un diplôme, un brevet » (Arrêté ministériel n° 1275 du 27/09/2022)", size=11, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, "", space_after=24)

    tbl = doc.add_table(rows=3, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    info = [
        ("Présenté par:", "[Nom(s) de l'étudiant(s)] — Matricule"),
        ("Encadré par:", "[Nom de l'encadrant] — Grade: [M.C.A / Pr]"),
        ("Année universitaire:", "2026 / 2027"),
    ]
    for i, (k, v) in enumerate(info):
        set_cell(tbl.rows[i].cells[0], k, bold=True)
        set_cell(tbl.rows[i].cells[1], v)
    page_break(doc)

    # Dedication
    heading(doc, "Dédicace", 1, center=True)
    para(doc, "[Dedication text — optional. Typically 1 paragraph. Remove this page if not needed.]", italic=True)
    page_break(doc)

    # Acknowledgments
    heading(doc, "Remerciements", 1, center=True)
    para(doc, "[Acknowledgments — thank your advisor, jury members, family. 1–2 paragraphs.]", italic=True)
    page_break(doc)

    # Abstract AR
    heading(doc, "الملخص", 1, center=True)
    para(doc, "[نص الملخص بالعربية — 150-250 كلمة: خلفية، مشكلة، أهداف، منهجية، أهم النتائج، كلمات مفتاحية.]", align=WD_ALIGN_PARAGRAPH.RIGHT)
    para(doc, "الكلمات المفتاحية: المحاسبة الآلية، النظام الضريبي الجزائري، المؤسسات الصغيرة والمتوسطة، الذكاء الاصطناعي، التحليل المالي.", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    page_break(doc)

    # Abstract EN
    heading(doc, "Abstract", 1, center=True)
    para(doc, "[English abstract 150–250 words: background, problem, objectives, methodology, key results — 35 screens, 20 ratios, Algerian tax engine, 1800 tests, 100% coverage.]")
    # Abstract FR
    heading(doc, "Résumé", 2)
    para(doc, "[Résumé en français — équivalent du résumé anglais.]")
    page_break(doc)

    # TOC (field)
    heading(doc, "Table of Contents", 1, center=True)
    p = doc.add_paragraph()
    run = p.add_run()
    set_run_font(run, size=11)
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin"); fld1.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click → Update Field to generate the table of contents."
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(sep); run._r.append(t); run._r.append(fld2)
    page_break(doc)

    # Lists
    heading(doc, "List of Figures", 1, center=True)
    para(doc, "[Figure 1.1  SME population growth 2012–2022 ................. 3]", space_after=3, line=1.2)
    para(doc, "[Figure 4.1  Layered architecture ............................ 21]", space_after=3, line=1.2)
    page_break(doc)
    heading(doc, "List of Tables", 1, center=True)
    para(doc, "[Table 2.2  The 20 implemented ratios ....................... 12]", space_after=3, line=1.2)
    page_break(doc)
    heading(doc, "List of Abbreviations", 1, center=True)
    for abbr in ["AIS", "IBS", "TVA", "IRG", "CNAS", "CNAC", "VF", "DAS", "RTL", "UAT", "ROE", "RSS"]:
        para(doc, f"[{abbr}] = [full name as used in thesis]", size=11, space_after=2, line=1.2)
    page_break(doc)

    # ========== Body section (decimal, restart at 1) ==========
    sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_page(doc, sec2)
    sec2.footer.is_linked_to_previous = False
    fp2 = sec2.footer.paragraphs[0]
    fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(fp2)
    set_page_number_format(sec2, "decimal", 1)

    for ch in range(1, 9):
        names = ["INTRODUCTION", "LITERATURE REVIEW", "PROBLEM ANALYSIS & REQUIREMENTS",
                 "PROPOSED SOLUTION (DESIGN)", "IMPLEMENTATION DETAILS",
                 "TESTING & RESULTS", "EVALUATION & DISCUSSION", "CONCLUSION & FUTURE WORK"]
        if ch > 1:
            page_break(doc)
        heading(doc, f"CHAPTER {roman(ch)}: {names[ch-1]}", 1, center=True)
        para(doc, f"Section {ch}.1 — Placeholder." , italic=True)
        heading(doc, f"{ch}.1 Section placeholder", 2)
        para(doc, "[Write the section content here. Follow the STYLE_GUIDE: Times New Roman 12, 1.5 line spacing, justified, APA citations. See thesis/OUTLINE.md for the detailed content plan of this chapter.]", italic=True)
        heading(doc, f"{ch}.2 Subsection placeholder", 3)
        para(doc, "[Subsection content.]", italic=True)
    page_break(doc)

    # Bibliography
    heading(doc, "BIBLIOGRAPHY", 1, center=True)
    para(doc, "References are listed in APA 7th edition. See thesis/REFERENCES_TEMPLATE.md for the master list and the copy-paste format blocks.", italic=True)
    for i in range(1, 9):
        para(doc, f"{i}. [Author, A. A. (Year). Title of work. Source.]", size=11, space_after=4, line=1.2)
    page_break(doc)

    # Appendices
    heading(doc, "APPENDICES", 1, center=True)
    for letter, title in [("A", "Installation & User Guide"),
                          ("B", "Screen Catalogue (35 screens)"),
                          ("C", "Tax Calculation Test Vectors"),
                          ("D", "Coverage Report Summary"),
                          ("E", "Innovation / Startup Dossier Notes (Arrêté 1275)")]:
        heading(doc, f"Appendix {letter} — {title}", 3)
        para(doc, "[Content placeholder — see OUTLINE.md Appendix register.]", italic=True)

    out = "thesis/thesis_template.docx"
    doc.save(out)
    print("SAVED:", out)


def roman(n):
    vals = {1: "I", 2: "II", 3: "III", 4: "IV", 5: "V", 6: "VI", 7: "VII", 8: "VIII"}
    return vals[n]


if __name__ == "__main__":
    build()
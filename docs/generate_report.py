# -*- coding: utf-8 -*-
"""مولّد تقرير المشروع (DOCX ثم PDF) بترقيم عربي صحيح من اليمين لليسار (RTL)."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DOCX = os.path.join(os.path.dirname(__file__), "Smart_Accounting_Platform_Report.docx")
FONT = "Tahoma"
MONO = "Consolas"
GREEN = RGBColor(0x1F, 0x7A, 0x33)
DARK = RGBColor(0x1E, 0x3A, 0x5F)
GRAY = RGBColor(0x44, 0x44, 0x44)


def _set_run_font(run, font=FONT, size=None, bold=None, italic=None, color=None):
    run.font.name = font
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _rtl_paragraph(p):
    ppr = p._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    ppr.append(bidi)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _rtl_run(run):
    rpr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rpr.append(rtl)


def ar_para(doc, text, size=11, bold=False, color=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    _rtl_paragraph(p)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold, color=color)
    _rtl_run(run)
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    _rtl_paragraph(p)
    sizes = {1: 16, 2: 13, 3: 12}
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    _set_run_font(run, size=sizes.get(level, 12), bold=True, color=GREEN if level == 1 else DARK)
    _rtl_run(run)
    if level == 1:
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F7A33")
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p


def bullet(doc, text, size=11, bold=False):
    p = doc.add_paragraph()
    _rtl_paragraph(p)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("• ")
    _set_run_font(run, size=size, bold=True)
    _rtl_run(run)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    _rtl_run(run)
    return p


def subbullet(doc, text, size=11):
    p = doc.add_paragraph()
    _rtl_paragraph(p)
    p.paragraph_format.left_indent = Cm(1.6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("◦ ")
    _set_run_font(run, size=size)
    _rtl_run(run)
    run = p.add_run(text)
    _set_run_font(run, size=size)
    _rtl_run(run)
    return p


def code_block(doc, code):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(code)
    _set_run_font(run, font=MONO, size=8.5, color=GRAY)
    return p


def set_cell(cell, text, bold=False, size=10, shade=None):
    p = cell.paragraphs[0]
    _rtl_paragraph(p)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    _set_run_font(run, size=size, bold=bold)
    _rtl_run(run)
    if shade:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), shade)
        tcPr.append(shd)


def make_table(doc, header, rows, widths=None, header_fill="1F7A33"):
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for i, h in enumerate(header):
        set_cell(t.rows[0].cells[i], h, bold=True, shade=header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], str(v))
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    return p


def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def build():
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)

    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)

    # ============================ غلاف ============================
    for _ in range(3):
        spacer(doc, 10)
    ar_para(doc, "الجمهورية الجزائرية الديمقراطية الشعبية", size=16, bold=True, space_after=4)
    ar_para(doc, "وزارة التعليم العالي والبحث العلمي", size=14, space_after=18)
    ar_para(doc, "مشروع تخرج", size=14, bold=True, color=GRAY, space_after=30)
    ar_para(doc, "المنصة المحاسبية الذكية", size=30, bold=True, color=GREEN, space_after=4)
    ar_para(doc, "Smart Accounting Platform", size=18, space_after=10)
    ar_para(doc, "نظام محاسبي متكامل مع الذكاء الاصطناعي", size=15, bold=True, space_after=40)
    ar_para(doc, "الإصدار: v3.1.7", size=14, bold=True, space_after=4)
    ar_para(doc, "السنة الجامعية: 2025/2026", size=13, space_after=8)
    page_break(doc)

    # ============================ الفهرس ============================
    heading(doc, "فهرس المحتويات", 1)
    toc_items = [
        "1. المقدمة",
        "2. مشكلة البحث",
        "3. أهداف المشروع",
        "4. أهمية المشروع",
        "5. الدراسات السابقة",
        "6. منهجية العمل",
        "7. الهندسة المعمارية للنظام",
        "8. التقنيات المستخدمة",
        "9. وصف الوحدات البرمجية",
        "10. واجهات المستخدم",
        "11. النظام الجبائي الجزائري",
        "12. نظام التحديث التلقائي",
        "13. الميزات المتقدمة الجديدة",
        "14. الاختبارات وضمان الجودة",
        "15. النتائج",
        "16. الخاتمة",
        "17. الأعمال المستقبلية",
    ]
    for item in toc_items:
        ar_para(doc, item, size=12, space_after=4)
    page_break(doc)

    # ============================ 1. المقدمة ============================
    heading(doc, "1. المقدمة", 1)
    ar_para(doc, "تعد المحاسبة المالية أحد أهم الركائز الأساسية لأي مؤسسة اقتصادية، حيث توفر المعلومات الضرورية لاتخاذ القرارات الاستراتيجية. مع التطور التكنولوجي المتسارع، أصبحت الأنظمة المحاسبية التقليدية غير قادرة على مواكبة احتياجات العصر الحديث من سرعة ودقة وتحليل متقدم.")
    ar_para(doc, "تهدف هذه المنصة المحاسبية الذكية إلى توفير حل متكامل يجمع بين أساسيات المحاسبة المالية وتقنيات الذكاء الاصطناعي وتحليل البيانات، مع التركيز على النظام الجبائي الجزائري كميزة أساسية تميزها عن الحلول الأخرى.")
    ar_para(doc, "تم تطوير المنصة باستخدام لغة Python وإطار العمل PyQt5، مع الاعتماد على أحدث التقنيات في تحليل البيانات والرسوم البيانية والتقارير. المشروع هو تطبيق سطح مكتب (Desktop Application) يعمل على نظام Windows 10/11، مع إمكانية تشغيله بشكل محمول (Portable) أو من خلال مثبت (Installer).")
    ar_para(doc, "وصل المشروع إلى الإصدار v3.1.7 بعد جلسات تطوير متتالية أضافت ميزات متقدمة شملت تحليل DuPont، تحليل السيناريوهات، المعايير المرجعية، لوحة التحكم المتقدمة، الامتثال الضريبي الكامل، محرك الرؤى الذكية، تحليل مراكز التكلفة، تعدد العملات، المزامنة السحابية، الشركات التجريبية، واختبار UAT شامل (1786 اختباراً ناجحاً).")

    # ============================ 2. مشكلة البحث ============================
    heading(doc, "2. مشكلة البحث", 1)
    ar_para(doc, "تعاني المؤسسات الصغيرة والمتوسطة في الجزائر من نقص في الحلول المحاسبية المتكاملة التي:")
    bullet(doc, "تتوافق مع النظام الجبائي الجزائري (IBS، TVA، IRG، CNAS، CNAC، DAS، G50/G57)")
    bullet(doc, "تقدم تحليلاً مالياً متقدماً ونسباً مالية دقيقة")
    bullet(doc, "توفر واجهة سهلة الاستخدام باللغة العربية مع دعم كامل للاتجاه من اليمين لليسار (RTL)")
    bullet(doc, "تدمج تقنيات الذكاء الاصطناعي للمساعدة في اتخاذ القرارات")
    bullet(doc, "تتيح التحديث التلقائي للبرنامج")
    bullet(doc, "تكون مفتوحة المصدر وقابلة للتطوير")
    ar_para(doc, "لذلك، جاءت فكرة هذا المشروع لسد هذه الفجوة وتوفير حل محاسبي متكامل يجمع بين الدقة والسهولة والتوافق مع البيئة الجزائرية.")

    # ============================ 3. أهداف المشروع ============================
    heading(doc, "3. أهداف المشروع", 1)
    bullet(doc, "أ. بناء نظام محاسبي متكامل يغطي جميع العمليات المالية الأساسية")
    bullet(doc, "ب. تطوير محرك تحليل مالي ينتج 20 نسبة مالية في 4 فئات إضافة إلى درجة Z-Score للتنبؤ بالإفلاس")
    bullet(doc, "ج. تطبيق النظام الجبائي الجزائري بشكل كامل (IBS، TVA 6%/9%/19%، IRG، CNAS، CNAC، VF، DAS)")
    bullet(doc, "د. توفير 28 شاشة تفاعلية بواجهة سهلة الاستخدام مع اختصارات لوحة مفاتيح كاملة")
    bullet(doc, "هـ. دعم 3 لغات: العربية، الإنجليزية، الفرنسية مع تبديل فوري (1530 مفتاحاً للترجمة)")
    bullet(doc, "و. دمج تقنيات الذكاء الاصطناعي (تنبؤات، كشف شذوذ، أنماط، توصيات ذكية)")
    bullet(doc, "ز. تطبيق نظام التحديث التلقائي (Auto-update) كامل بتحديث مخفي تماماً")
    bullet(doc, "ح. توفير 1786 اختبار وحدة ناجحاً مع تغطية 100% للوحدات البرمجية")
    bullet(doc, "ط. توفير تقارير متعددة الصيغ (PDF/Excel/HTML/TXT) بدعم كامل للعربية")
    bullet(doc, "ي. دعم 3 ثيمات (فاتح، داكن، عصري) مع تحسينات تباين وإتاحة")

    # ============================ 4. أهمية المشروع ============================
    heading(doc, "4. أهمية المشروع", 1)
    ar_para(doc, "تتمثل أهمية هذا المشروع في كونه أول منصة محاسبية مفتوحة المصدر تستهدف المؤسسات الجزائرية بشكل خاص، مع مراعاة الخصوصية الثقافية واللغوية والجبائية. كما أنه يوفر أدوات تحليل متقدمة تساعد المحاسبين والمديرين الماليين في اتخاذ قرارات استراتيجية مبنية على بيانات دقيقة.")
    ar_para(doc, "بالإضافة إلى ذلك، يعتبر المشروع مرجعاً تعليمياً متكاملاً يجمع بين المحاسبة المالية، تحليل البيانات، الذكاء الاصطناعي، وهندسة البرمجيات. حيث يمكن للطلاب تعلم كيفية بناء نظام محاسبي متكامل من الصفر باستخدام أحدث التقنيات.")

    # ============================ 5. الدراسات السابقة ============================
    heading(doc, "5. الدراسات السابقة", 1)
    ar_para(doc, "تمت مراجعة العديد من الأنظمة المحاسبية المتوفرة في السوق ودراستها، ومن أبرزها: برامج المحاسبة التجارية الجاهزة، الأنظمة المحاسبية السحابية، وأدوات التحليل المالي المنفصلة. لوحظ أن معظم هذه الحلول إما عامة لا تراعي خصوصية البيئة الجزائرية، أو غير متوافقة مع النظام الجبائي المحلي.")
    ar_para(doc, "يتميز هذا المشروع عن سابقاته بدمج النظام الجبائي الجزائري بشكل كامل، وتوفير تحليل مالي متقدم مع 20 نسبة مالية ودرجة Z-Score، ودعم 3 لغات، ونظام تحديث تلقائي. كما أنه مفتوح المصدر بالكامل (رخصة MIT)، مما يسمح لأي مطور بتعديله وتطويره.")

    # ============================ 6. منهجية العمل ============================
    heading(doc, "6. منهجية العمل", 1)
    ar_para(doc, "تم اعتماد منهجية التطوير التكراري (Iterative Development) مع اتباع بروتوكول التعديل الجراحي (Surgical Editing Protocol).")
    ar_para(doc, "مبادئ المنهجية:", bold=True)
    bullet(doc, "المس فقط ما يجب لمسه: لا تحسين تنسيق كود مجاور، لا إعادة صياغة تعليقات قديمة")
    bullet(doc, "مطابقة الأسلوب: الالتزام بأسلوب الكود الحالي تماماً حتى لو كان غير مثالي")
    bullet(doc, "تحليل التأثير: قراءة PROJECT_MAP.md، تحديد الملفات المتأثرة بدقة قبل كل تغيير")
    bullet(doc, "السلامة المعمارية والتجريد: الالتزام بـ DRY، استخدام طبقة Shared/Core")
    bullet(doc, "التحقق والنجاح: كتابة الاختبار أولاً (TDD)، تأكيد فشله، ثم جعله ينجح")
    bullet(doc, "مزامنة الحالة: تحديث PROJECT_MAP.md فوراً بعد كل تعديل")
    ar_para(doc, "مراحل التطوير:", bold=True)
    bullet(doc, "1. تحليل المتطلبات وجمع المعلومات")
    bullet(doc, "2. تصميم قاعدة البيانات والهيكلة (أكثر من 60 جدولاً مع تطور المشروع)")
    bullet(doc, "3. تطوير الوحدات الأساسية (calculations.py، analysis.py، audit.py)")
    bullet(doc, "4. تطوير واجهات المستخدم (28 شاشة)")
    bullet(doc, "5. تطبيق النظام الجبائي الجزائري (tax.py + tax_config.json)")
    bullet(doc, "6. إضافة تقنيات الذكاء الاصطناعي (تنبؤات + كشف شذوذ + توصيات)")
    bullet(doc, "7. الاختبار والتحسين المستمر (1786 اختباراً + تغطية 100%)")
    bullet(doc, "8. النشر والتوزيع (Nuitka + Inno Setup + GitHub)")

    # ============================ 7. الهندسة المعمارية ============================
    heading(doc, "7. الهندسة المعمارية للنظام", 1)
    ar_para(doc, "تم بناء النظام وفق نمط العمارة متعددة الطبقات (Layered Architecture) لفصل الاهتمامات وتسهيل الصيانة والتطوير.")
    heading(doc, "7.1 الطبقات", 2)
    bullet(doc, "طبقة العرض (UI): 28 شاشة عبر PyQt5 مع تحميل كسول (Lazy Loading)")
    bullet(doc, "طبقة المنطق التجاري (Modules): أكثر من 25 وحدة برمجية متخصصة")
    bullet(doc, "طبقة البيانات (Database): SQLite مع تجمّع اتصالات (Connection Pool)")
    bullet(doc, "طبقة الأدوات المشتركة (Utils): تنسيق، تحقق، أمان")
    heading(doc, "7.2 تدفق البيانات", 2)
    ar_para(doc, "التدفق العام للبيانات في النظام:")
    code_block(doc, "واجهة الإدخال (DataEntryView)\n    \u2193\nAppState (حالة التطبيق عبر pyqtSignal)\n    \u2193\nCalculationEngine (حساب 20 نسبة + Z-Score)\n    \u2193\nAnalysisEngine (تحليل DuPont + اتجاهات + سيناريوهات)\n    \u2193\nAuditEngine (8 فحوصات تدقيق)\n    \u2193\nTaxEngine (الجباية الجزائرية الكاملة)\n    \u2193\nReportingEngine (PDF/Excel/HTML/TXT)\n    \u2193\nDashboardView (4 رسوم بيانية)")

    # ============================ 8. التقنيات المستخدمة ============================
    heading(doc, "8. التقنيات المستخدمة", 1)
    ar_para(doc, "تم اختيار التقنيات بناءً على معايير: الأداء، سهولة التطوير، التوافق مع النظام، المجتمع النشط، والتكلفة (جميعها مفتوحة المصدر ومجانية).")
    heading(doc, "8.1 لغة Python", 2)
    ar_para(doc, "تم اختيار Python لعدة أسباب: سهولة التعلم والقراءة، مجتمع كبير ونشط، مكتبات غنية لتحليل البيانات والرسوم البيانية، والقدرة على تحويل الكود إلى تطبيق مستقل (Standalone) باستخدام Nuitka.")
    heading(doc, "8.2 PyQt5 — إطار عمل واجهات المستخدم", 2)
    ar_para(doc, "يتم إنشاء MainWindow (QMainWindow) عند بدء التشغيل. تحتوي على QStackedWidget يضم 28 شاشة (QWidget). التنقل بين الشاشات يتم عبر QListWidget جانبي. يتم تحميل الشاشات عند الطلب (Lazy Loading) لتحسين الأداء.")
    code_block(doc, "def _switch_view(self, index):\n    widget = self._get_or_create_view(index)\n    self.stacked_widget.setCurrentWidget(widget)\n    self.sidebar_list.setCurrentRow(index - 1)")
    heading(doc, "8.3 SQLite — قاعدة البيانات", 2)
    ar_para(doc, "يتم إنشاء قاعدة البيانات SQLite عند أول تشغيل للتطبيق. تم اعتماد تجمّع اتصالات (Connection Pool) في db_connection.py مع دفع عمليات الإدراج (executemany) في db_operations.py مما حسّن سرعة الحفظ 4.6 مرة والقراءة 17 مرة.")
    heading(doc, "8.4 Pandas — تحليل البيانات", 2)
    ar_para(doc, "مكتبة Python الرائدة لتحليل البيانات. تستخدم لحساب النسب المالية، تحليل الاتجاهات، كشف الشذوذ (z-score + IQR)، التنبؤ (الانحدار الخطي/المتوسط المتحرك/التجانس الأسي)، وتحضير البيانات للرسوم البيانية. استُخدمت pandas وnumpy فقط دون sklearn لتقليل الاعتماديات.")
    heading(doc, "8.5 Matplotlib — الرسوم البيانية", 2)
    ar_para(doc, "مكتبة الرسوم البيانية الأساسية. تُستخدم في لوحة التحكم، الرادار، شلال DuPont، حساسية Tornado، خطوط التنبؤ، ومخططات مراكز التكلفة. يتم دمجها مع PyQt5 عبر FigureCanvasQTAgg وتحديثها تلقائياً عند تغير البيانات.")
    heading(doc, "8.6 FPDF2 — تصدير PDF بالعربية", 2)
    ar_para(doc, "مكتبة متخصصة في إنشاء ملفات PDF. تم اختيارها لقدرتها على دعم الخطوط العربية مع دعم الكتابة من اليمين لليسار (RTL). يتم تحميل خط Amiri كخط مضمّن (Embedded Font) في ملفات PDF المصدرة.")
    heading(doc, "8.7 Nuitka — تحويل Python إلى C/EXE", 2)
    ar_para(doc, "مترجم Python يحول الكود إلى C ثم يجمعه إلى ملف تنفيذي. أنتج نسخة مستقلة (Standalone EXE) حجمها 133MB. تم تحسين الإقلاع من 778ms إلى 49ms وتقليل استهلاك الذاكرة من 128MB إلى 45MB.")
    heading(doc, "8.8 Inno Setup — إنشاء المثبت", 2)
    ar_para(doc, "أداة إنشاء مثبتات Windows. توزّع التطبيق بمثبت حجمه 63.2MB ونسخة ZIP بحجم 41.4MB. الميزات: AppMutex لضمان إغلاق التطبيق، CloseApplications=force، وتحديث صامت.")
    heading(doc, "8.9 Cryptography — التشفير والأمان", 2)
    ar_para(doc, "تستخدم لتشفير كلمات المرور، توليد رموز المصادقة الثنائية (2FA)، تشفير النسخ الاحتياطية السحابية بكلمة مرور عبر AES-GCM، والتحقق من سلامة البيانات.")
    heading(doc, "8.10 OpenPyXL — تصدير Excel", 2)
    ar_para(doc, "مكتبة للتعامل مع ملفات Excel (.xlsx) لتصدير التقارير بتنسيق احترافي (ألوان، حدود، عرض أعمدة) مع دعم كل شاشات التقارير المتقدمة.")

    # ============================ 9. وصف الوحدات البرمجية ============================
    heading(doc, "9. وصف الوحدات البرمجية", 1)
    heading(doc, "9.1 محرك الحسابات (modules/calculations.py)", 2)
    ar_para(doc, "【الموقع: 20 نسبة - 4 فئات + Z-Score】")
    ar_para(doc, "يحسب 20 نسبة مالية موزعة على 4 فئات رئيسية: السيولة، الربحية، الكفاءة، المديونية. إضافة إلى درجة Z-Score للتنبؤ باحتمالية الإفلاس، و3 حقول إدخال جديدة (النقدية، المصاريف التشغيلية، متوسط الموردين) تمكن من حساب نسب أكثر دقة مثل Cash Ratio، فترة المخزون، وفترة سداد الموردين.")
    heading(doc, "9.2 محرك التحليل (modules/analysis.py)", 2)
    ar_para(doc, "【الموقع: DuPont متقدم - اتجاهات - تدفقات - سيناريوهات】")
    bullet(doc, "تحليل DuPont (3 و5 عوامل) مع مؤشر ROE وشلال توضيحي")
    bullet(doc, "تحليل رأس المال العامل والتدفقات النقدية (تشغيلية/استثمارية/تمويلية)")
    bullet(doc, "تحليل الاتجاهات عبر السنوات ومقارنة مع الميزانية")
    bullet(doc, "تحليل السيناريوهات: مثالي/طبيعي/أسوأ مع حساسية Tornado")
    heading(doc, "9.3 محرك التدقيق (modules/audit.py)", 2)
    ar_para(doc, "【الموقع: 8 فحوصات تدقيق مالي آلية】")
    ar_para(doc, "يتحقق من صحة معادلة الميزانية، قائمة الدخل، التدفقات النقدية، تحليل الانحرافات، فحص نسب السيولة والربحية والمديونية، وفحص الاستمرارية. كل فحص يرجع كائن AuditResult يحتوي على الحالة (pass/fail/warning) والرسالة التوضيحية.")
    heading(doc, "9.4 النظام الجبائي الجزائري (modules/tax.py)", 2)
    ar_para(doc, "【الموقع: IBS - TVA - IRG - CNAS - CNAC - VF - DAS】")
    bullet(doc, "IBS: 19% للإنتاج، 26% للتجارة والخدمات، مع الحد الأدنى")
    bullet(doc, "TVA: 19% قياسي، 9% مخفض، 6% للمواد الوسيطة (ميزة جديدة)، 0% للتصدير")
    bullet(doc, "IRG: شرائح تصاعدية من 0% إلى 35%")
    bullet(doc, "CNAS: 11% و CNAC: 1% (لا تتضاعف مع عدد الموظفين)")
    bullet(doc, "VF: 0.5% من رقم الأعمال")
    bullet(doc, "DAS: بيانات الأجور السنوية مع قوالب الإقرارات")
    ar_para(doc, "تتم قراءة المعدلات من tax_config.json لتسهيل التحديث.")
    heading(doc, "9.5 التقارير الضريبية (modules/tax_reports.py)", 2)
    ar_para(doc, "【الموقع: G50 - G57 - DAS - IBS】")
    ar_para(doc, "ينشئ قوالب الإقرارات الضريبية G50 (TVA) وG57 (IBS) وإقرار DAS (البيان السنوي للأجور) مع تصدير PDF وExcel، إضافة إلى ترحيل رصيد TVA ودفعات IBS المقدمة في الأشهر 3 و6 و11 وتصفية IBS.")
    heading(doc, "9.6 المعايير المرجعية (modules/benchmarks.py)", 2)
    ar_para(doc, "【الموقع: 7 قطاعات - 10 نسب - اتجاه - منافسون】")
    ar_para(doc, "مقارنة الأداء المالي مع معايير 7 قطاعات صناعية (الصناعات التحويلية، التجارة والتوزيع، الخدمات، البناء، الفلاحة، النقل واللوجستيك، التكنولوجيا). تشمل أفضل الممارسات والمعيار الدولي ونقاط القوة/الضعف والاتجاه عبر السنوات ومقارنة المنافسين (جدولا reference_standards وcompetitor_data).")
    heading(doc, "9.7 محرك الرؤى الذكية (modules/ai_insights.py)", 2)
    ar_para(doc, "【الموقع: تنبؤ - شذوذ - أنماط - توصيات】")
    bullet(doc, "التنبؤ 3-6 أشهر: انحدار خطي / متوسط متحرك / تجانس أسي مع فترات ثقة 95%")
    bullet(doc, "كشف الشذوذ: z-score للأرباح + IQR للمعاملات")
    bullet(doc, "أنماط: اتجاه/موسمية/دورات/مؤشرات مخاطر")
    bullet(doc, "توصيات ذكية وتنبيهات (خطر/تحذير/فرصة/إجراء) مع تصدير PDF/Excel")
    ar_para(doc, "يعتمد على pandas وnumpy فقط دون sklearn.")
    heading(doc, "9.8 ربحية مراكز التكلفة (modules/cost_center_profitability.py)", 2)
    ar_para(doc, "【الموقع: مراكز - توزيع - ربحية - اتجاه】")
    ar_para(doc, "يعرّف مراكز التكلفة (قسم/مشروع/فرع/خط إنتاج)، يوزع التكاليف المباشرة وغير المباشرة (إيرادات/عدد موظفين/مساحة/متساوٍ)، ويحلل الربحية مع مقارنات (سابقة/ميزانية/معايير) واتجاه متعدد الفترات وتوصيات وتصدير PDF/Excel.")
    heading(doc, "9.9 لوحة التحكم المتقدمة (modules/advanced_dashboard.py)", 2)
    ar_para(doc, "【الموقع: 6 KPI - 4 رسوم - تنبيهات - تخصيص】")
    ar_para(doc, "6 بطاقات KPI بحالة لونية، 4 رسوم (إيرادات شهرية/ربعية + مصروفات + ربحية + رادار)، تنبيهات (شذوذ/أداء/معايير/إجراءات)، تخصيص كامل مع حفظ التخطيطات في جدول dashboard_layouts، وتصدير PDF/Excel.")
    heading(doc, "9.10 وحدات أخرى", 2)
    bullet(doc, "modules/reporting.py: تقارير TXT/HTML/PDF(عربي)/Excel")
    bullet(doc, "modules/user_manager.py: 4 أدوار - 16 صلاحية - 2FA - استعادة كلمة المرور")
    bullet(doc, "modules/update_checker.py: فحص وتحميل وتثبيت صامت وإعادة تشغيل")
    bullet(doc, "modules/currency.py: 7 عملات افتراضية + أسعار صرف + تحويل + تقرير متعدد العملات")
    bullet(doc, "modules/cloud_sync.py: مزامنة (Dropbox/OneDrive/Drive) + snapshot مع checksum + تشفير AES-GCM + نسخ احتياطي مع تدوير")
    bullet(doc, "modules/demo_data.py: 4 شركات تجريبية بمولّد معاملات شهرية (12 شهراً بأوزان موسمية)")
    bullet(doc, "modules/cashflow.py: قائمة التدفقات النقدية (تشغيلية/استثمارية/تمويلية)")
    bullet(doc, "modules/comparative.py: تحليل مقارن متعدد السنوات/الشركات")
    bullet(doc, "modules/breakeven.py: تحليل نقطة التعادل")
    bullet(doc, "modules/forecasting.py: التنبؤ المالي")
    bullet(doc, "modules/backup.py + scheduled_backup.py: نسخ احتياطي واسترجاع")
    bullet(doc, "modules/email_notifier.py: إشعارات البريد")
    bullet(doc, "modules/print_manager.py: طباعة التقارير")
    bullet(doc, "modules/data_import.py + csv_import.py: استيراد البيانات من Excel/CSV")

    # ============================ 10. واجهات المستخدم ============================
    heading(doc, "10. واجهات المستخدم", 1)
    ar_para(doc, "يحتوي النظام على 28 شاشة تفاعلية بواجهة سهلة الاستخدام، مع دعم كامل للغة العربية من اليمين لليسار (RTL).")
    heading(doc, "10.1 المميزات العامة", 2)
    bullet(doc, "دعم 3 لغات (عربي RTL، إنجليزي LTR، فرنسي LTR) مع تبديل فوري (1530 مفتاح ترجمة)")
    bullet(doc, "3 ثيمات: فاتح (Light)، داكن (Dark)، عصري (Modern) مع تحسينات تباين وإتاحة")
    bullet(doc, "اختصارات كاملة: Ctrl+1..0 و Ctrl+Shift+1..0 و F2..F8 للشاشات 21-28 و Ctrl+T للثيم و F1 للاختصارات")
    bullet(doc, "قائمة عرض (View) ديناميكية بكل الشاشات 28")
    bullet(doc, "انتقالات تلاشي عند تغيير الشاشة ومؤشر تحميل في شريط الحالة")
    bullet(doc, "رسائل خطأ موحّدة مع إجراء مقترح مترجم (ui/widgets/messages.py)")
    bullet(doc, "تحميل كسول (Lazy Loading) لتحسين الأداء")
    heading(doc, "10.2 الشاشات الـ28", 2)
    make_table(doc,
        ["#", "الشاشة", "#", "الشاشة"],
        [
            ["1", "لوحة التحكم", "15", "المعايير المرجعية"],
            ["2", "إدخال البيانات", "16", "مخطط الحسابات"],
            ["3", "النسب المالية (20 نسبة + Z-Score)", "17", "البنك"],
            ["4", "التقويم الجبائي", "18", "التقارير"],
            ["5", "الطباعة", "19", "النسخ الاحتياطي"],
            ["6", "تحليل DuPont المتقدم", "20", "الاستيراد"],
            ["7", "تحليل السيناريوهات", "21", "التنبؤ المالي"],
            ["8", "لوحة التحكم المتقدمة", "22", "مراكز التكلفة"],
            ["9", "محرك الرؤى الذكية AI", "23", "تعدد العملات"],
            ["10", "مراكز التكلفة الربحية", "24", "المزامنة السحابية"],
            ["11", "الشركات التجريبية", "25", "المستخدمون"],
            ["12", "التقارير الضريبية", "26", "الإعدادات"],
            ["13", "البيانات المالية", "27", "مركز المساعدة"],
            ["14", "المقارنات المالية", "28", "المعاينة والتصدير"],
        ],
        widths=[1.2, 7.0, 1.2, 7.0])
    spacer(doc, 4)

    # ============================ 11. النظام الجبائي الجزائري ============================
    heading(doc, "11. النظام الجبائي الجزائري", 1)
    ar_para(doc, "يعتبر النظام الجبائي الجزائري من أكثر الأنظمة تعقيداً في المنطقة، حيث يجمع بين ضرائب متعددة ومعدلات متفاوتة حسب النشاط والقطاع. تم تطوير هذا الموديول ليشمل جميع الجوانب الأساسية للجباية الجزائرية.")
    heading(doc, "11.1 IBS (ضريبة أرباح الشركات)", 2)
    bullet(doc, "النشاط الإنتاجي (صناعة، فلاحة): 19%")
    bullet(doc, "النشاط التجاري والخدمات: 26%")
    bullet(doc, "الحد الأدنى: 500,000 دج أو 0.5% من رقم الأعمال للشركات ذات رقم أعمال أكبر من 5,000,000 دج")
    bullet(doc, "الشركات الجديدة: إعفاء لمدة 3 سنوات")
    heading(doc, "11.2 TVA (ضريبة القيمة المضافة)", 2)
    bullet(doc, "المعدل القياسي: 19% (معظم السلع والخدمات)")
    bullet(doc, "المعدل المخفض: 9% (المواد الغذائية الأساسية، الكهرباء، الغاز)")
    bullet(doc, "المعدل الوسيط: 6% (المواد الوسيطة — ميزة جديدة)")
    bullet(doc, "المعدل الصفري: 0% (التصدير)")
    bullet(doc, "ترحيل رصيد TVA وإقرار G50 الشهري")
    heading(doc, "11.3 IRG (ضريبة الدخل الإجمالي)", 2)
    ar_para(doc, "ضريبة تصاعدية على الدخل الفردي بشرائح تبدأ من 0% (إعفاء حتى 30,000 دج شهرياً) وصولاً إلى 35% لأعلى الشريحة، مع خصم 50% من الدخل الخاضع للضريبة.")
    heading(doc, "11.4 CNAS و CNAC", 2)
    bullet(doc, "CNAS: 11% (7% على صاحب العمل، 4% على العامل)")
    bullet(doc, "CNAC: 1% (صندوق البطالة — يتحملها العامل بالكامل)")
    bullet(doc, "ملاحظة: لا تتضاعف الاشتراكات مع عدد الموظفين (أُصلح في محاكاة الضرائب)")
    heading(doc, "11.5 VF (الرسم المقطوع)", 2)
    ar_para(doc, "ضريبة على الشركات غير المنتظمة في تصاريحها، بمعدل 0.5% من رقم الأعمال وتُضاف إلى IBS.")
    heading(doc, "11.6 DAS و الإقرارات", 2)
    bullet(doc, "DAS: البيان السنوي للأجور (بيانات DAS + تصدير PDF/Excel)")
    bullet(doc, "G50: إقرار TVA الشهري/الربع سنوي")
    bullet(doc, "G57: إقرار IBS")
    bullet(doc, "دفعات IBS المقدمة في الأشهر 3 و6 و11 مع تصفية نهاية السنة")
    heading(doc, "11.7 المحاكاة الجبائية الكاملة", 2)
    ar_para(doc, "تقوم TaxView بمحاكاة جبائية كاملة: إدخال البيانات، حساب جميع الضرائب تلقائياً، عرض ملخص الجباية، تحليل العبء الجبائي، وإنشاء تقارير قابلة للتصدير.")

    # ============================ 12. نظام التحديث التلقائي ============================
    heading(doc, "12. نظام التحديث التلقائي", 1)
    ar_para(doc, "تم تطوير نظام تحديث تلقائي مستوحى من نظام تحديث opencode. يسمح النظام بتحديث التطبيق بنقرة واحدة دون الحاجة لفتح المتصفح.")
    heading(doc, "12.1 آلية العمل", 2)
    bullet(doc, "1. عند بدء التشغيل: خيط منفصل يفحص التحديثات عبر GitHub Pages (docs/version.json)")
    bullet(doc, "2. مقارنة الإصدار المحلي مع الإصدار البعيد")
    bullet(doc, "3. عرض QMessageBox مع قائمة التغييرات (changelog)")
    bullet(doc, "4. تحميل المثبت عبر QProgressDialog")
    bullet(doc, "5. التثبيت بـ /SILENT (صامت)")
    bullet(doc, "6. إعادة التشغيل عبر wscript/VBS مخفي تماماً (بدون نافذة cmd)")
    heading(doc, "12.2 معالجة الأخطاء", 2)
    bullet(doc, "فشل الاتصال بالإنترنت: تسجيل الخطأ في السجل دون إزعاج المستخدم")
    bullet(doc, "فشل التحميل: إغلاق التطبيق دون تغيير الإصدار الحالي")
    bullet(doc, "تعارض Mutex: إغلاق التطبيق عبر CloseApplications=force في Inno Setup")

    # ============================ 13. الميزات المتقدمة الجديدة ============================
    heading(doc, "13. الميزات المتقدمة الجديدة", 1)
    ar_para(doc, "شملت جلسات التطوير الأخيرة (من v3.1.3 إلى v3.1.7) إضافة مجموعة من الميزات المتقدمة التي تُعدّ جوهر الابتكار في هذا المشروع:")

    heading(doc, "13.1 شاشة تحليل DuPont المتقدمة", 2)
    bullet(doc, "شلال (Waterfall) + خط (Trend) + مؤشر ROE مع مقارنة قطاعية")
    bullet(doc, "توصيات تلقائية بناءً على نتائج العوامل الثلاثة")
    bullet(doc, "تصدير PDF/Excel (modules/analysis.py + analysis_view.py)")

    heading(doc, "13.2 شاشة تحليل السيناريوهات", 2)
    bullet(doc, "3 سيناريوهات: مثالي/طبيعي/أسوأ مع حساسية Tornado")
    bullet(doc, "رسوم خط/شريط/مساحة مع تصدير PDF")
    bullet(doc, "حفظ السيناريوهات JSON أو في قاعدة البيانات")

    heading(doc, "13.3 المعايير المرجعية المتقدمة", 2)
    bullet(doc, "أفضل الممارسات + معيار دولي + نقاط قوة/ضعف")
    bullet(doc, "اتجاه عبر السنوات + مقارنة منافسين (جدولا reference_standards وcompetitor_data)")
    bullet(doc, "مخطط الرادار (Radar Chart) للقطاعات")

    heading(doc, "13.4 لوحة التحكم المتقدمة", 2)
    bullet(doc, "6 بطاقات KPI بحالة لونية (أخضر/برتقالي/أحمر)")
    bullet(doc, "4 رسوم: إيرادات شهرية/ربعية + مصروفات + ربحية + رادار")
    bullet(doc, "تنبيهات: شذوذ/أداء/معايير/إجراءات")
    bullet(doc, "تخصيص كامل + حفظ التخطيطات في جدول dashboard_layouts + تصدير PDF/Excel")

    heading(doc, "13.5 الامتثال الضريبي الجزائري", 2)
    bullet(doc, "نسبة TVA 6% (intermediate) + ترحيل رصيد TVA")
    bullet(doc, "دفعات IBS المقدمة (أشهر 3، 6، 11) + تصفية IBS")
    bullet(doc, "بيانات DAS + قوالب إقرارات G50/G57/DAS مع تصدير PDF/Excel")

    heading(doc, "13.6 محرك الرؤى الذكية AI (شاشة 24)", 2)
    bullet(doc, "تنبؤ 3-6 أشهر (خطي/متوسط متحرك/تجانس أسي + فترات ثقة 95%)")
    bullet(doc, "كشف شذوذ (z-score للأرباح + IQR للمعاملات)")
    bullet(doc, "أنماط (اتجاه/موسمية/دورات/مؤشرات مخاطر)")
    bullet(doc, "توصيات ذكية + تنبيهات (خطر/تحذير/فرصة/إجراء) + تصدير PDF/Excel")

    heading(doc, "13.7 تحليل ربحية مراكز التكلفة (شاشة 25)", 2)
    bullet(doc, "مراكز: قسم/مشروع/فرع/خط إنتاج")
    bullet(doc, "توزيع مباشر/غير مباشر (إيرادات/عدد موظفين/مساحة/متساوٍ)")
    bullet(doc, "مقارنات (سابقة/ميزانية/معايير) + اتجاه متعدد الفترات + تصدير PDF/Excel")

    heading(doc, "13.8 تعدد العملات (شاشة 26)", 2)
    bullet(doc, "7 عملات افتراضية + أسعار صرف + محول")
    bullet(doc, "تقرير متعدد العملات + تصدير CSV (modules/currency.py)")

    heading(doc, "13.9 المزامنة السحابية والنسخ الاحتياطي (شاشة 27)", 2)
    bullet(doc, "وجهات مزامنة (Dropbox/OneDrive/Drive)")
    bullet(doc, "snapshot مع checksum + تشفير اختياري بكلمة مرور (AES-GCM)")
    bullet(doc, "نسخ احتياطي تلقائي مع تدوير + استرجاع/سحب + سجل عمليات DB")

    heading(doc, "13.10 الشركات التجريبية (شاشة 28)", 2)
    bullet(doc, "4 شركات (تجارية/خدمات/إنتاج/استيراد-تصدير) ببيانات مالية وضريبية متسقة")
    bullet(doc, "مولّد معاملات شهرية (12 شهراً بأوزان موسمية)")
    bullet(doc, "تقارير مُعدّة مسبقاً + تصدير CSV + قوالب استيراد متوافقة مع DataImporter")

    heading(doc, "13.11 تحسين الأداء الشامل (v3.1.6)", 2)
    bullet(doc, "تحميل كسول للمشاهد (main_window.py) + PEP 562 للوحدات (modules/__init__.py)")
    bullet(doc, "تجمّع اتصالات DB (db_connection.py + close_pool) + دفعات executemany")
    bullet(doc, "إقلاع 778→49ms، ذاكرة 128→45MB، حفظ DB 4.6×، قراءة 17×، لوحة تحكم بدون إعادة رسم (التفاصيل في docs/PERFORMANCE_REPORT.md)")
    heading(doc, "13.12 اختبار UAT الشامل (v3.1.7)", 2)
    bullet(doc, "رحلة مستخدم حقيقية عبر التطبيق: تسجيل دخول + تجوّل 35 شاشة + إدخال بيانات + تبديل 3 لغات + حفظ + تسجيل خروج (tests/test_uat.py)")
    bullet(doc, "كشف وإصلاح بغّ حقيقي: تبديل اللغة على شاشتي المزامنة البنكية واستيراد البيانات كان يسقط التطبيق (_clear_layout)")
    bullet(doc, "1786 اختباراً ناجحاً + تغطية وحدات 100%")

    heading(doc, "13.12 تحسينات UI/UX", 2)
    bullet(doc, "اختصارات كاملة (Ctrl+1..0، Ctrl+Shift+1..0، F2..F8، Ctrl+T، F1)")
    bullet(doc, "قائمة عرض ديناميكية + انتقالات تلاشي + مؤشر تحميل")
    bullet(doc, "رسائل خطأ موحّدة مع إجراء مقترح مترجم + تحسينات إتاحة/تباين في الثيمات الثلاثة")

    # ============================ 14. الاختبارات وضمان الجودة ============================
    heading(doc, "14. الاختبارات وضمان الجودة", 1)
    ar_para(doc, "تم تطوير 1229 اختبار وحدة موزعة على 45 ملف اختبار، تغطي جميع مكونات النظام، عبر `python -m pytest tests -q` (و1090 عبر unittest).")
    heading(doc, "14.1 تغطية الوحدات", 2)
    ar_para(doc, "بلغت تغطية الوحدات البرمجية 100% عبر `python -m coverage run --source=modules --omit=\"modules/__init__.py\" -m pytest tests -q` (كانت 73% ثم 99%)، وكل ملفات modules/ عند 100% (5768 سطراً بلا أي سطر مفقود).")
    heading(doc, "14.2 ملفات الاختبار الجديدة", 2)
    make_table(doc,
        ["ملف الاختبار", "الوحدات المغطاة", "ملف الاختبار", "الوحدات المغطاة"],
        [
            ["test_edge_errors", "edge/error للميزات الست", "test_importers", "data_import + csv_import"],
            ["test_cashflow", "cashflow", "test_excel_export", "excel_export"],
            ["test_comparative", "comparative", "test_reporting_extra", "reporting"],
            ["test_backup", "backup + scheduled_backup", "test_user_manager", "user_manager"],
            ["test_bank_print", "bank_sync + print_manager", "test_update_checker_extra", "update_checker"],
            ["test_breakeven_costcenter", "breakeven + cost_center + forecasting", "test_cloud_sync_extra", "cloud_sync"],
            ["test_reporting_modules", "calculations + report_templates + activity_log", "test_tax_reminders_extra", "tax_reminders"],
            ["test_email_currency", "email_notifier + currency", "test_tax_reports_extra", "tax_reports"],
            ["test_small_gaps", "budget + validation + advanced_dashboard", "الملفات القديمة", "الميزات الأساسية"],
        ],
        widths=[5.2, 6.0, 5.2, 6.0])
    spacer(doc, 4)
    heading(doc, "14.3 ضمان الجودة", 2)
    bullet(doc, "pytest 9.1.1: إطار الاختبار الأساسي")
    bullet(doc, "coverage 7.15.2: قياس تغطية الأكواد")
    bullet(doc, "منهجية TDD: كتابة الاختبار ثم جعله ينجح")
    bullet(doc, "i18n: 1530 مفتاحاً × 3 لغات (AR/EN/FR)")
    ar_para(doc, "أُصلح خلال جلسة التغطية خلل واحد: modules/comparative.py generate_report كان يرمي KeyError مع بيانات ناقصة واستُبدل بـ .get(item/ratio, 0). كما وُثّقت قائمة أخطاء كاملة في PROJECT_MAP.md.")

    # ============================ 15. النتائج ============================
    heading(doc, "15. النتائج", 1)
    ar_para(doc, "تم تحقيق الأهداف التالية:" )
    bullet(doc, "28 شاشة تفاعلية كاملة مع دعم RTL للعربية و3 لغات و3 ثيمات")
    bullet(doc, "20 نسبة مالية + Z-Score + تحليل DuPont متقدم + سيناريوهات")
    bullet(doc, "نظام جبائي جزائري كامل (IBS، TVA 6/9/19%، IRG، CNAS، CNAC، VF، DAS، G50/G57)")
    bullet(doc, "محرك رؤى ذكية بالتنبؤ وكشف الشذوذ والتوصيات")
    bullet(doc, "1786 اختباراً ناجحاً مع تغطية 100% للوحدات")
    bullet(doc, "أداء محسّن: إقلاع 49ms وذاكرة 45MB")
    bullet(doc, "مثبت Inno Setup (51.6MB) ونسخة ZIP (82.9MB) وتحديث تلقائي مخفي")

    # ============================ 16. الخاتمة ============================
    heading(doc, "16. الخاتمة", 1)
    ar_para(doc, "تم في هذا المشروع بناء منصة محاسبية ذكية متكاملة تجمع بين أحدث تقنيات البرمجة وتحليل البيانات المالية، مع التركيز على النظام الجبائي الجزائري كميزة أساسية. يقدم المشروع حلاً عملياً للمؤسسات الصغيرة والمتوسطة في الجزائر، ويوفر أدوات تحليل متقدمة تساعد في اتخاذ القرارات المالية.")
    ar_para(doc, "تم تطوير المشروع وفق أفضل ممارسات هندسة البرمجيات، مع الالتزام بمبادئ الجودة والاختبار المستمر، مما نتج عنه 1786 اختبار وحدة ناجحاً وتغطية 100% للوحدات. كما تم توفير نظام تحديث تلقائي لتسهيل عملية التطوير والصيانة المستمرة.")
    ar_para(doc, "المشروع مفتوح المصدر بالكامل (رخصة MIT) على GitHub، مما يسمح بالمساهمة والتطوير من قبل المجتمع.")

    # ============================ 17. الأعمال المستقبلية ============================
    heading(doc, "17. الأعمال المستقبلية", 1)
    heading(doc, "17.1 المدى القصير", 2)
    bullet(doc, "1. ربط API مع البنوك الجزائرية للمزامنة البنكية التلقائية")
    bullet(doc, "2. دعم الفواتير الإلكترونية (E-Invoicing)")
    bullet(doc, "3. توقيع رقمي للنسخ التنفيذية وتجاوز Smart App Control")
    heading(doc, "17.2 المدى المتوسط", 2)
    bullet(doc, "1. تطوير نسخة سحابية (Cloud) مع قاعدة بيانات مركزية")
    bullet(doc, "2. إضافة وحدات إدارة المخزون والمشتريات")
    bullet(doc, "3. إضافة تقارير مالية معيارية (IFRS، IAS)")
    bullet(doc, "4. تحليل متقدم باستخدام Machine Learning للتنبؤ بالأزمات المالية")
    heading(doc, "17.3 المدى البعيد", 2)
    bullet(doc, "1. تطوير تطبيق جوال (Mobile App) لمتابعة المؤشرات")
    bullet(doc, "2. ربط مع مصلحة الضرائب الجزائرية للتصريح الإلكتروني")
    bullet(doc, "3. دعم التوقيع الإلكتروني")
    bullet(doc, "4. منصة ذكاء اصطناعي متكاملة للتحليل المالي والتنبؤ")

    doc.save(OUT_DOCX)
    print("SAVED:", OUT_DOCX)


if __name__ == "__main__":
    build()
